import sys
import requests
import json

from docx import *
from tkinter import *
from tkinter import messagebox



def parse():
	document = Document('hierarchal content.docx')
	copy =Document()
	for paragraph in document.paragraphs:

		if paragraph.style.name =='Heading 1':
			copy.add_paragraph().add_run(paragraph.text).bold =True

		elif paragraph.style.name =='Heading 2':
			copy.add_paragraph().add_run(paragraph.text).italic =True

		else:
			copy.add_paragraph().add_run(paragraph.text)

	copy.save('new document.docx')
	messagebox.showinfo("message", 'File Parsed')


def store():
	#sever endpoint
	URL='http://localhost:8000/reqs/'

	document = Document('hierarchal content.docx')

	headNum=0
	paraNum=0

	for paragraph in document.paragraphs:	
		if paragraph.style.name =='Heading 1' or paragraph.style.name =='Heading 2':
			headNum+=1
			
		else:
			paraNum+=1

	params={'headNum':headNum,
			'paraNum':paraNum
			}
	
	#making request
	try:
		r = requests.post(url=URL,json=params)
		json_data = json.loads(r.text)
		messagebox.showinfo("message", json_data['message'])
	except Exception:
		messagebox.showinfo("error", 'no response')




def visualize():
	#endpoint
	URL='http://localhost:8000/reqs/'

	try:
		r =requests.get(url=URL)
		res = json.loads(r.text)
		
		json_data = json.loads(r.text)

		headNum=int(float(json_data['headNum']))
		paraNum=int(float(json_data['paraNum']))

		document = Document('template.docx')
		copy = Document()

		for paragraph in document.paragraphs:
			
			if paragraph.text == '#_numberofHeadlines#':
				copy.add_paragraph().add_run(str(headNum))

			elif paragraph.text == '#_numberofParagraphs#':
				copy.add_paragraph().add_run(str(paraNum))

			else:
				copy.add_paragraph().add_run(paragraph.text)

		copy.save('template.docx')
		messagebox.showinfo("message", 'Statistics saved')

	except Exception as e:
		messagebox.showinfo("error", e)
	


def main():
	app.title("Python Developer Test")
	yPos=60
	xPos=200

	f = Frame(app,bg="orange",width=400,height=270)
	f.grid(row=0,column=0,sticky="NW")
	f.grid_propagate(0)
	f.update()
	
	label = Label(f,text="Hi! We've been expecting you",bg="yellow",fg="red",font=("Arial", 17))
	label.place(x=xPos, y=yPos, anchor="center")
	
	
	btn1 = Button(f, text="Parse",bg="yellow",fg="red",font=("Arial", 15),command=parse)
	btn1.place(x=xPos, y=yPos+50,anchor="center")
	
	
	btn2 = Button(f, text="Store Statistics",bg="yellow",fg="red",font=("Arial", 15),command=store)
	btn2.place(x=xPos, y=yPos+100,anchor="center") 
	
	
	btn3 = Button(f, text="Visualize Statistics",bg="yellow",fg="red",font=("Arial", 15),command=visualize)
	btn3.place(x=xPos, y=yPos+150,anchor="center") 

	app.bind('<Escape>', close)
	app.geometry('400x270')
	app.resizable(width=False, height=False)
	app.mainloop()


def close(event):
    app.withdraw() 
    sys.exit() 



if __name__ == '__main__':
	app = Tk() 
	main()


